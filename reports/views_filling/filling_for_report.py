from .filling_summary_widget import summarry_widget_image
from .filling_volume_widget import volume_widget_image
from .filling_top_10_authors_by_volume_widget import top_10_authors_by_volume_widget_image
from .filling_top_10_sources_widget import top_10_sources_widget_image
from .filling_top_10_countries_widget import top_10_countries_widget_image
from .filling_top_10_languages_widget import top_10_languages_widget_image
from .filling_sentiment_top_10_sources_widget import sentiment_top_10_sources_widget_image
from .filling_sentiment_top_10_authors_widget import sentiment_top_10_authors_widget_image
from .filling_sentiment_top_10_countries_widget import sentiment_top_10_countries_widget_image
from .filling_sentiment_top_10_languages_widget import sentiment_top_10_languages_widget_image
from .filling_sentiment_for_period_widget import sentiment_for_period_widget_image
from .filling_content_volume_top_5_source_widget import content_volume_top_5_source_widget_image
from .filling_content_volume_top_5_authors_widget import content_volume_top_5_authors_widget_image
from .filling_content_volume_top_5_countries_widget import content_volume_top_5_countries_widget_image
from docx.shared import Pt, Inches
from .options import *
from django.shortcuts import get_object_or_404
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor
from project.models import Project
from widgets.models import WidgetsList2, ClippingFeedContentWidget

def export_title(project_id):
  proj = get_object_or_404(Project, pk=project_id)
  title = proj.title
  return title

def export_period(project_id):
  proj = get_object_or_404(Project, pk=project_id)
  start_d = str(proj.start_search_date.ctime())
  end_d = str(proj.end_search_date.ctime())
  period = start_d + ' - ' + end_d
  return period

def font_one(run, cell):
  font = run.font
  font.bold = True
  font.size = Pt(15)
  para = cell.add_paragraph()
  para.paragraph_format.line_spacing = Inches(0.1)

def font_two(run, cell):
  font = run.font
  font.size = Pt(11) 
  para = cell.add_paragraph()
  para.paragraph_format.line_spacing = Inches(0.1)

def filling_templates_for_instant_and_regular_reports(document, project_id):
  tbls = document.tables
  for t in tbls:
    for row in t.rows:
      for cell in row.cells:
        for p in cell.paragraphs:
          if '$EXPORT_TITLE$' in p.text:
            p.text = p.text.replace('$EXPORT_TITLE$', export_title(project_id))
          if '$EXPORT_PERIOD$' in p.text:
            p.text = p.text.replace('$EXPORT_PERIOD$', export_period(project_id))
          if '$EXPORT_TOC$' in p.text:
            proj = WidgetsList2.objects.get(project_id=project_id)
            p.text = p.text.replace('$EXPORT_TOC$',' ')
            if proj.summary.is_active:
              run = cell.add_paragraph('').add_run('Report Summary')
              font_one(run, cell)
              run = cell.add_paragraph().add_run('Summary')
              font_two(run, cell)
              cell.add_paragraph()
            if proj.volume.is_active or proj.top_authors.is_active or proj.top_brands.is_active or proj.top_countries.is_active or proj.top_languages.is_active:
              run = cell.add_paragraph('').add_run('Potential Reach')
              font_one(run, cell)
              if proj.volume.is_active:
                run = cell.add_paragraph().add_run(proj.volume.title + ' (per ' + proj.volume.aggregation_period + ')')
                font_two(run, cell)
              if proj.top_authors.is_active:
                run = cell.add_paragraph().add_run(proj.top_authors.title + ' (per ' + proj.top_authors.aggregation_period + ')')
                font_two(run, cell)
              if proj.top_brands.is_active:
                run = cell.add_paragraph().add_run(proj.top_brands.title + ' (per ' + proj.top_brands.aggregation_period + ')')
                font_two(run, cell)
              if proj.top_countries.is_active:
                run = cell.add_paragraph().add_run(proj.top_countries.title + ' (per ' + proj.top_countries.aggregation_period + ')')
                font_two(run, cell)
              if proj.top_languages.is_active:
                run = cell.add_paragraph().add_run(proj.top_languages.title + ' (per ' + proj.top_languages.aggregation_period + ')')
                font_two(run, cell)
              cell.add_paragraph()  
            if proj.sentiment_top_sources.is_active or proj.sentiment_top_authors.is_active or proj.sentiment_top_countries.is_active or proj.sentiment_top_languages.is_active or proj.sentiment_for_period.is_active:
              run = cell.add_paragraph().add_run('Sentiment Distribution')
              font_one(run, cell)
              if proj.sentiment_top_sources.is_active:
                run = cell.add_paragraph().add_run(proj.sentiment_top_sources.title)
                font_two(run, cell)
              if proj.sentiment_top_authors.is_active:
                run = cell.add_paragraph().add_run(proj.sentiment_top_authors.title)
                font_two(run, cell)
              if proj.sentiment_top_countries.is_active:
                run = cell.add_paragraph().add_run(proj.sentiment_top_countries.title)
                font_two(run, cell)
              if proj.sentiment_top_languages.is_active:
                run = cell.add_paragraph().add_run(proj.sentiment_top_languages.title)
                font_two(run, cell)
              if proj.sentiment_for_period.is_active:
                run = cell.add_paragraph().add_run(proj.sentiment_for_period.title)
                font_two(run, cell)
                cell.add_paragraph()
            if proj.content_volume_top_sources.is_active or proj.content_volume_top_authors.is_active:
              run = cell.add_paragraph().add_run('Content Volume')
              font_one(run, cell)
              if proj.content_volume_top_sources.is_active:
                run = cell.add_paragraph().add_run(proj.content_volume_top_sources.title)
                font_two(run, cell)
              if proj.content_volume_top_authors.is_active:
                run = cell.add_paragraph().add_run(proj.content_volume_top_authors.title)
                font_two(run, cell)
              if proj.content_volume_top_countries.is_active:
                run = cell.add_paragraph().add_run(proj.content_volume_top_countries.title)
                font_two(run, cell)
              cell.add_paragraph()  
            if proj.clipping_feed_content.is_active:
              run = cell.add_paragraph().add_run('Clippings: Sample Articles')
              font_one(run, cell)

  def new_section(name_section):
    p3 = document.add_paragraph()
    p = p3._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '055FFC')
    pBdr.append(bottom)
    p3 = p3.add_run(name_section)
    p3.font.size = Pt(20)
    document.add_paragraph()

  if proj.summary.is_active:
    new_section('Report Summary')
    summarry_widget_image(document, proj)
    document.add_page_break()

  if proj.volume.is_active or proj.content_volume_top_sources.is_active or proj.content_volume_top_authors.is_active or proj.content_volume_top_countries.is_active:
    new_section('Content Volume')
    volume_widget_image(document, proj)
    content_volume_top_5_source_widget_image(document, proj)
    content_volume_top_5_authors_widget_image(document, proj)
    content_volume_top_5_countries_widget_image(document, proj)
    document.add_page_break()
  
  if proj.top_authors.is_active or proj.top_brands.is_active or proj.top_countries.is_active or proj.top_languages.is_active:
    new_section('Top 10')
    top_10_authors_by_volume_widget_image(document, proj)
    top_10_sources_widget_image(document, proj)
    top_10_countries_widget_image(document, proj)
    top_10_languages_widget_image(document, proj)
    document.add_page_break()

  if proj.volume.is_active or proj.top_authors.is_active or proj.top_brands.is_active or proj.top_countries.is_active or proj.top_languages.is_active:
    new_section('Sentiment Top 10')
    sentiment_top_10_sources_widget_image(document, proj)
    sentiment_top_10_authors_widget_image(document, proj)
    sentiment_top_10_countries_widget_image(document, proj)
    sentiment_top_10_languages_widget_image(document, proj)
    document.add_paragraph()
    document.add_paragraph()
    sentiment_for_period_widget_image(document, proj)
    document.add_page_break()

  if proj.clipping_feed_content.is_active:
    new_section('Clippings: Sample Articles')
    posts = ClippingFeedContentWidget.objects.filter(project_id=project_id).values(
        'post__entry_title',
        'post__feed_language__language',
        'post__feedlink__alexaglobalrank',
        'post__sentiment',
        'post__feedlink__source1',
        'post__feedlink__country',
        'post__entry_published',
        'post__entry_summary',
        )
    for each in posts:
      table = document.add_table(rows=6, cols=5)
      a = table.cell(0, 0)
      b = table.cell(0, 4)
      title = a.merge(b)
      a = table.cell(1, 0)
      b = table.cell(1, 4)
      content = a.merge(b)
      flag = 0
      for row in table.rows:
        if flag == 0:
          row.height = Inches(0.5)
          for cell in row.cells:
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), "#e9eaeb")
            cell._tc.get_or_add_tcPr().append(shade_obj)# set silver background color
        if flag == 1:
          row.height = Inches(1.1)
          for cell in row.cells:
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), "#ffffff")
            cell._tc.get_or_add_tcPr().append(shade_obj)
        if flag == 2 or flag == 4:
          row.height = Inches(0.5)
          for cell in row.cells:
              shade_obj = OxmlElement('w:shd')
              shade_obj.set(qn('w:fill'), "#f0f1f2")
              cell._tc.get_or_add_tcPr().append(shade_obj)# set silver background color
        if flag == 3 or flag == 5:
          row.height = Inches(0.5)
          for cell in row.cells:
              shade_obj = OxmlElement('w:shd')
              shade_obj.set(qn('w:fill'), "#ffffff")
              cell._tc.get_or_add_tcPr().append(shade_obj)
        flag += 1
      title.text = each['post__entry_title'] if each['post__entry_title'] else ''
      title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
      title.paragraphs[0].runs[0].font.size = Pt(12)
      def font_three(cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(77, 77, 77)
      def font_four(cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(95, 95, 95)  
      content.text = each['post__entry_summary'] if each['post__entry_summary'] else ''
      font_three(content)
      heading_cells = table.rows[2].cells
      heading_cells[0].text = 'Media Channel'
      font_four(heading_cells[0])
      heading_cells[1].text = 'Language'
      font_four(heading_cells[1])
      heading_cells[2].text = 'Potential reach'
      font_four(heading_cells[2])
      heading_cells[3].text = 'Sentiment'
      font_four(heading_cells[3])
      heading_cells[4].text = 'Source'
      font_four(heading_cells[4])
      sec_row = table.rows[3].cells
      sec_row[0].text = 'Online'
      font_three(sec_row[0])
      sec_row[1].text = each['post__feed_language__language'] if each['post__feed_language__language'] else ''
      font_three(sec_row[1])
      sec_row[2].text = str(each['post__feedlink__alexaglobalrank']) if each['post__feedlink__alexaglobalrank'] else ''
      font_three(sec_row[2])
      sec_row[3].text = each['post__sentiment'] if each['post__sentiment'] else ''
      font_three(sec_row[3])
      sec_row[4].text = each['post__feedlink__source1'] if each['post__feedlink__source1'] else ''
      font_three(sec_row[4])
      heading_cells = table.rows[4].cells
      heading_cells[0].text = 'Source country'
      font_four(heading_cells[0])
      heading_cells[1].text = 'Domestic Rank'
      font_four(heading_cells[1])
      heading_cells[2].text = 'Global Rank'
      font_four(heading_cells[2])
      heading_cells[3].text = 'Date'
      font_four(heading_cells[3])
      sec_row = table.rows[5].cells
      sec_row[0].text = each['post__feedlink__country'] if each['post__feedlink__country'] else ''
      font_three(sec_row[0])
      sec_row[1].text = '1'
      font_three(sec_row[1])
      sec_row[2].text = '1'
      font_three(sec_row[2])
      sec_row[3].text = each['post__entry_published'].strftime("%d.%m.%Y")
      font_three(sec_row[3])
      document.add_paragraph()
      document.add_paragraph()
  return document
