from docx.shared import Inches, Pt

def widget_image(document, title_widget, widget_image):
    document.add_paragraph()
    p = document.add_paragraph().add_run(title_widget)
    p.font.size = Pt(13)
    document.add_paragraph()
    document.add_picture(widget_image, width=Inches(6.6))
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
