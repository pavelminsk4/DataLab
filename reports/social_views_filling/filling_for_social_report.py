from .filling_social_summary_widget import social_summarry_widget_image
from .filling_social_top_locations_widget import social_top_locations_widget_image
from docx.shared import Pt, Inches
from django.shortcuts import get_object_or_404
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor
from project_social.models import *

def export_title(project_id):
  proj = get_object_or_404(ProjectSocial, pk=project_id)
  title = proj.title
  return title

def export_period(project_id):
  proj = get_object_or_404(ProjectSocial, pk=project_id)
  start_d = str(proj.start_search_date.ctime())
  end_d = str(proj.end_search_date.ctime())
  period = start_d + ' - ' + end_d
  return period

def font_one(run, cell): # Need to move to separate file from filling_for_report also
  font = run.font
  font.bold = True
  font.size = Pt(15)
  para = cell.add_paragraph()
  para.paragraph_format.line_spacing = Inches(0.1)

def font_two(run, cell):
  font = run.font
  font.size = Pt(11) 
  para = cell.add_paragraph()
  para.paragraph_format.line_spacing = Inches(0.1)

def filling_templates_for_social_reports(document, item, screenshots_list):
  proj = ProjectSocial.objects.get(id=item.module_project_id)
  tbls = document.tables
  for t in tbls:
    for row in t.rows:
      for cell in row.cells:
        for p in cell.paragraphs:
          if '$EXPORT_TITLE$' in p.text:
            p.text = p.text.replace('$EXPORT_TITLE$', export_title(item.module_project_id))
          if '$EXPORT_PERIOD$' in p.text:
            p.text = p.text.replace('$EXPORT_PERIOD$', export_period(item.module_project_id))
          if '$EXPORT_TOC$' in p.text:
            p.text = p.text.replace('$EXPORT_TOC$',' ')
            if item.soc_summary:
              run = cell.add_paragraph('').add_run('Report Summary')
              font_one(run, cell)
              run = cell.add_paragraph().add_run('Summary')
              font_two(run, cell)
              cell.add_paragraph()
            if item.soc_content_volume or item.soc_content_volume_by_top_locations or item.soc_content_volume_by_top_authors or item.soc_content_volume_by_top_languages:
              run = cell.add_paragraph('').add_run('Potential Reach')
              font_one(run, cell)
              if item.soc_content_volume:
                run = cell.add_paragraph().add_run(proj.social_widgets_list.content_volume.title + ' (per ' + proj.social_widgets_list.content_volume.aggregation_period + ')')
                font_two(run, cell)
              if item.soc_content_volume_by_top_locations:
                run = cell.add_paragraph().add_run(proj.social_widgets_list.content_volume_by_top_locations.title + ' (per ' + proj.social_widgets_list.content_volume_by_top_locations.aggregation_period + ')')
                font_two(run, cell)
              if item.soc_content_volume_by_top_authors:
                run = cell.add_paragraph().add_run(proj.social_widgets_list.content_volume_by_top_authors.title + ' (per ' + proj.social_widgets_list.content_volume_by_top_authors.aggregation_period + ')')
                font_two(run, cell)
              if item.soc_content_volume_by_top_languages:
                run = cell.add_paragraph().add_run(proj.social_widgets_list.content_volume_by_top_languages.title + ' (per ' + proj.social_widgets_list.content_volume_by_top_languages.aggregation_period + ')')
                font_two(run, cell)
              cell.add_paragraph()  

  def new_section(name_section):
    p3 = document.add_paragraph()
    p = p3._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '055FFC')
    pBdr.append(bottom)
    p3 = p3.add_run(name_section)
    p3.font.size = Pt(20)
    document.add_paragraph()

  if item.soc_summary:
    social_summarry_widget_image(document, screenshots_list['summary'])
    document.add_page_break()

  if item.soc_top_locations:
    social_top_locations_widget_image(document, screenshots_list['top_locations'])
    document.add_page_break()

  return document
