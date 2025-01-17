from project_social.models import ProjectSocial
from docx.shared import Pt, Inches


class ReportDocument:
    def __init__(self, document, item, screenshot_list, module, widget_list):
        self.document = document
        self.item = item
        self.screenshot_list = screenshot_list
        self.project = module.objects.get(id=item.module_project_id)
        self.widget_list = getattr(self.project, widget_list)

    def fill(self):
        self.__foreach_paragraph(self.__fill_introduction)
        self.__fill_content()
        return self.document

    def __font_one(self, run, cell):
        font = run.font
        font.bold = True
        font.size = Pt(15)
        para = cell.add_paragraph()
        para.paragraph_format.line_spacing = Inches(0.1)

    def __font_two(self, run, cell):
        font = run.font
        font.size = Pt(11)
        para = cell.add_paragraph()
        para.paragraph_format.line_spacing = Inches(0.1)

    def __export_period(self):
        start_d = str(self.project.start_search_date.ctime())
        end_d = str(self.project.end_search_date.ctime())
        period = f'{start_d} - {end_d}'
        return period

    def __foreach_paragraph(self, func):
        tbls = self.document.tables
        for t in tbls:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        func(p, cell)

    def __fill_titular(self, p, cell):
        if '$EXPORT_TITLE$' in p.text:
            p.text = p.text.replace('$EXPORT_TITLE$', self.project.title)
        if '$EXPORT_PERIOD$' in p.text:
            p.text = p.text.replace('$EXPORT_PERIOD$', self.__export_period())

    def __fill_table_of_contents(self, p, cell):
        if '$EXPORT_TOC$' in p.text:
            p.text = p.text.replace('$EXPORT_TOC$', ' ')
            self.__fill_summary_section(p, cell)
            self.__fill_sentiment_section(p, cell)
            self.__fill_demography_section(p, cell)
            self.__fill_influencers_section(p, cell)

    def __fill_summary_section(self, p, cell):
        widgets = {
                    'soc_summary': ['summary', False],
                    'soc_sentiment': ['sentiment', False],
                    'soc_top_locations': ['top_locations', False],
                    'soc_top_authors': ['top_authors', False],
                    'soc_top_languages': ['top_languages', False],
                    'soc_content_volume': ['content_volume', True],
                    'soc_content_volume_top_locations': ['content_volume_top_locations', True],
                    'soc_content_volume_top_authors': ['content_volume_top_authors', True],
                    'soc_content_volume_top_languages': ['content_volume_top_languages', True],
                    'soc_gender_volume': ['gender_volume', True],
                    
                    'onl_summary': ['summary', False],
                    'onl_volume': ['volume', True],
                    'onl_sentiment_for_period': ['sentiment_for_period', True],
                    'onl_top_sources': ['top_sources', False],
                    'onl_top_authors': ['top_authors', False],
                    'onl_top_keywords': ['top_keywords', False],
                    'onl_top_countries': ['top_countries', False],
                    'onl_top_languages': ['top_languages', False],
                    'onl_content_volume_top_sources': ['content_volume_top_sources', False],
                    'onl_content_volume_top_authors': ['content_volume_top_authors', False],
                    'onl_content_volume_top_countries': ['content_volume_top_countries', False],
                  }
        if True in [(True if getattr(self.item, widget) else False) for widget in widgets.keys()]:
            self.__font_one(cell.add_paragraph('').add_run('Summary'), cell)
        for widget in widgets.keys():
            if getattr(self.item, widget):
                self.__font_two(self.__get_widget_title(cell, getattr(self.widget_list, widgets[widget][0]), widgets[widget][1]), cell)
        cell.add_paragraph()

    def __fill_sentiment_section(self, p, cell):
        widgets = {
                    'soc_sentiment_diagram': ['sentiment_diagram', False],
                    'soc_sentiment_number_of_results': ['sentiment_number_of_results', False], 
                    'soc_sentiment_authors': ['sentiment_authors', False], 
                    'soc_sentiment_locations': ['sentiment_locations', False], 
                    'soc_sentiment_languages': ['sentiment_languages', False],
                    'soc_sentiment_by_gender': ['sentiment_by_gender', False],
                    'soc_sentiment_top_keywords': ['sentiment_top_keywords', False],
                    
                    'onl_sentiment_number_of_results': ['top_keywords', False],
                    'onl_sentiment_diagram': ['sentiment_diagram', False],
                    'onl_sentiment_top_sources': ['sentiment_top_sources', False],
                    'onl_sentiment_top_countries': ['sentiment_top_countries', False],
                    'onl_sentiment_top_authors': ['sentiment_top_authors', False],
                    'onl_sentiment_top_languages': ['sentiment_top_languages', False],
                    'onl_sentiment_top_keywords': ['sentiment_top_keywords', False],
                  }
        if True in [(True if getattr(self.item, widget) else False) for widget in widgets.keys()]:
            self.__font_one(cell.add_paragraph('').add_run('Sentiment'), cell)
        for widget in widgets.keys():
            if getattr(self.item, widget):
                self.__font_two(self.__get_widget_title(cell, getattr(self.widget_list, widgets[widget][0]), widgets[widget][1]), cell)
        cell.add_paragraph()
        
    def __fill_demography_section(self, p, cell):
        widgets = {
                    'soc_top_keywords': ['top_keywords', False],
                    'soc_authors_by_gender': ['authors_by_gender', False],
                    'soc_authors_by_language': ['authors_by_language', False],
                    'soc_authors_by_location': ['authors_by_location', False],
                    'soc_gender_by_location': ['gender_by_location', False],
                    'soc_keywords_by_location': ['keywords_by_location', False],
                    'soc_languages_by_location': ['languages_by_location', False],
                    
                    'onl_sources_by_country': ['sources_by_country', False],
                    'onl_authors_by_country': ['authors_by_country', False],
                    'onl_languages_by_country': ['languages_by_country', False],
                    'onl_keywords_by_country': ['top_keywords_by_country', False],
                  }
        if True in [(True if getattr(self.item, widget) else False) for widget in widgets.keys()]:
            self.__font_one(cell.add_paragraph('').add_run('Demography'), cell)
        for widget in widgets.keys():
            if getattr(self.item, widget):
                self.__font_two(self.__get_widget_title(cell, getattr(self.widget_list, widgets[widget][0]), widgets[widget][1]), cell)
        cell.add_paragraph()

    def __fill_influencers_section(self, p, cell):
        widgets = {
                    'soc_top_sharing_sources': ['top_sharing_sources', False],
                    'soc_authors_by_sentiment': ['authors_by_sentiment', False],
                    
                    'onl_top_sharing_sources': ['top_sharing_sources', False],
                    'onl_authors_by_language': ['authors_by_language', False],
                    'onl_overall_top_sources': ['overall_top_sources', False],
                    'onl_overall_top_authors': ['overall_top_authors', False],
                    'onl_authors_by_sentiment': ['authors_by_sentiment', False],
                    'onl_sources_by_language': ['sources_by_language', False],
                  }
        if True in [(True if getattr(self.item, widget) else False) for widget in widgets.keys()]:
            self.__font_one(cell.add_paragraph('').add_run('Influencers'), cell)
        for widget in widgets.keys():
            if getattr(self.item, widget):
                self.__font_two(self.__get_widget_title(cell, getattr(self.widget_list, widgets[widget][0]), widgets[widget][1]), cell)
        cell.add_paragraph()

    def __get_widget_title(self, cell, widget, have_period):
        return cell.add_paragraph().add_run(f"{widget.title}{f' (per {widget.aggregation_period})' if have_period else ''}")

    def __fill_content(self):
        widgets = {
                    'soc_summary': 'Summary',
                    'soc_sentiment': 'Sentiment',
                    'soc_top_locations': 'Top Locations',
                    'soc_top_authors': 'Top Authors',
                    'soc_top_languages': 'Top Languages',
                    'soc_content_volume': 'Content Volume',
                    'soc_content_volume_top_locations': 'Content Volume Top Locations',
                    'soc_content_volume_top_authors': 'Content Volume Top Authors',
                    'soc_content_volume_top_languages': 'Content Volume Top Languages',
                    'soc_gender_volume': 'Gender volume',
                    'soc_sentiment_diagram': 'Sentiment diagram',
                    'soc_sentiment_number_of_results': 'Sentiment number of results', 
                    'soc_sentiment_authors': 'Sentiment authors',
                    'soc_sentiment_locations': 'Sentiment locations',
                    'soc_sentiment_languages': 'Sentiment languages',
                    'soc_sentiment_by_gender': 'Sentiment by gender',
                    'soc_top_keywords': 'Top keywords',
                    'soc_sentiment_top_keywords': 'Sentiment top keywords',
                    'soc_authors_by_gender': 'Authors by gender',
                    'soc_authors_by_language': 'Authors by language',
                    'soc_authors_by_location': 'Authors by location',
                    'soc_gender_by_location': 'Top gender by location',
                    'soc_keywords_by_location': 'Top keywords by location',
                    'soc_languages_by_location': 'Top languages by location',
                    'soc_authors_by_sentiment': 'Authors by sentiment',
                    'soc_top_sharing_sources': 'Top sharing sources',
                    
                    'onl_summary': 'Summary',
                    'onl_volume': 'Content volume',
                    'onl_sentiment_for_period': 'Sentiment for period',
                    'onl_top_sources': 'Top sources',
                    'onl_top_authors': 'Top authors',
                    'onl_top_keywords': 'Top keywords',
                    'onl_top_countries': 'Top countries',
                    'onl_top_languages': 'Top languages',
                    'onl_content_volume_top_sources': 'Content Volume by top sources',
                    'onl_content_volume_top_authors': 'Content volume by top authors',
                    'onl_content_volume_top_countries': 'Content volume by top countries',
                    'onl_sentiment_number_of_results': 'Sentiment number of results',
                    'onl_sentiment_diagram': 'Sentiment diagram',
                    'onl_sentiment_top_sources': 'Sentiment top sources',
                    'onl_sentiment_top_countries': 'Sentiment top countries',
                    'onl_sentiment_top_authors': 'Sentiment top authors',
                    'onl_sentiment_top_languages': 'Sentiment top languages',
                    'onl_sentiment_top_keywords': 'Sentiment top keywords',
                    'onl_sources_by_country': 'Sources by country',
                    'onl_authors_by_country': 'Authors by country',
                    'onl_languages_by_country': 'Top languages by country',
                    'onl_keywords_by_country': 'Top keywords by country',
                    'onl_top_sharing_sources': 'Top sharing sources',
                    'onl_authors_by_language': 'Authors by language',
                    'onl_overall_top_sources': 'Overall top sources',
                    'onl_overall_top_authors': 'Overall top authors',
                    'onl_authors_by_sentiment': 'Authors by sentiment',
                    'onl_sources_by_language': 'Sources by language',
                  }
        for widget in widgets.keys():
            if getattr(self.item, widget):
                self.__get_widget_image(self.screenshot_list[widget], widgets[widget])

    def __get_widget_image(self, widget_image, title_widget):
        # self.document.add_paragraph()
        # p = self.document.add_paragraph().add_run(title_widget)
        # p.font.size = Pt(13)
        # self.document.add_paragraph()
        self.document.add_picture(widget_image, width=Inches(6.6))
        self.document.add_paragraph()
        self.document.add_paragraph()
        self.document.add_paragraph()

    def __fill_introduction(self, p, cell):
        self.__fill_titular(p, cell)
        self.__fill_table_of_contents(p, cell)
